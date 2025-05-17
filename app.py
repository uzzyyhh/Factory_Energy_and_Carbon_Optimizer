import streamlit as st
import pandas as pd
import numpy as np
import plotly.graph_objects as go
import plotly.express as px
from sklearn.ensemble import RandomForestRegressor
from sklearn.linear_model import LinearRegression
from sklearn.preprocessing import OneHotEncoder
from sklearn.compose import ColumnTransformer
from sklearn.pipeline import Pipeline
from sklearn.metrics import mean_absolute_error, mean_squared_error, r2_score
from sklearn.model_selection import train_test_split
import joblib
import logging
from datetime import datetime
from docx import Document
from io import BytesIO
import uuid

# Configure logging
logging.basicConfig(level=logging.INFO, filename='app.log')
logger = logging.getLogger(__name__)

# Constants
GRID_EMISSION_FACTOR = 0.5  # kg CO2/kWh for grid
SOLAR_EMISSION_FACTOR = 0.05  # kg CO2/kWh for solar
TREES_PER_TON_CO2 = 48  # Trees per ton of CO2 absorbed
MAX_ENERGY_INCREASE = 1.1  # Limit optimized energy to 10% above baseline

def process_uploaded_data(file):
    """Process uploaded Excel/CSV file and map to internal format."""
    try:
        if file.name.endswith('.xlsx'):
            df = pd.read_excel(file, engine='openpyxl')
        elif file.name.endswith('.csv'):
            df = pd.read_csv(file)
        else:
            raise ValueError("Unsupported file format. Please upload .xlsx or .csv.")

        required_cols = ['Timestamp', 'Total Consumption (kWh)', 'Machine 1 (kWh)', 'Machine 2 (kWh)', 'HVAC (kWh)', 'Lighting (kWh)', 'Other (kWh)']
        if not all(col in df.columns for col in required_cols):
            raise ValueError(f"Missing required columns. Expected: {required_cols}")

        df['Timestamp'] = pd.to_datetime(df['Timestamp'], format='%m/%d/%y %H:%M')
        df = df.rename(columns={
            'Total Consumption (kWh)': 'Total_Energy',
            'Machine 1 (kWh)': 'Energy_Heavy',
            'Machine 2 (kWh)': 'Energy_Medium',
            'HVAC (kWh)': 'HVAC_Energy',
            'Lighting (kWh)': 'Lighting_Energy',
            'Other (kWh)': 'Other_Energy'
        })

        df['Day_of_Week'] = df['Timestamp'].dt.dayofweek
        df['Hour'] = df['Timestamp'].dt.hour
        df['Is_Weekday'] = df['Day_of_Week'] < 5
        df['Is_Working_Hours'] = df['Is_Weekday'] & (df['Hour'] >= 8) & (df['Hour'] < 18)

        conditions = [
            (~df['Is_Weekday']),
            (df['Is_Weekday'] & (df['Hour'] >= 8) & (df['Hour'] < 16)),
            (df['Is_Weekday'] & (df['Hour'] >= 16) & (df['Hour'] < 24))
        ]
        choices = ['weekend', 'day', 'night']
        df['Shift'] = np.select(conditions, choices, default='overnight')

        max_heavy_energy = df['Energy_Heavy'].max()
        max_medium_energy = df['Energy_Medium'].max()
        df['Intended_Heavy_On'] = ((df['Energy_Heavy'] / max_heavy_energy) * 5).round().clip(0, 5)
        df['Intended_Heavy_On'] = np.where(df['Is_Working_Hours'], df['Intended_Heavy_On'], 0)
        df['Intended_Medium_On'] = ((df['Energy_Medium'] / max_medium_energy) * 10).round().clip(0, 10)
        df['Intended_Medium_On'] = np.where(df['Is_Working_Hours'], df['Intended_Medium_On'], 0)

        df['Heavy_On'] = df['Intended_Heavy_On'].copy()
        df['Medium_On'] = df['Intended_Medium_On'].copy()

        df['Temperature'] = 30 + 5 * np.sin(2 * np.pi * (df['Hour'] - 14) / 24) + np.random.normal(0, 1, len(df))
        df['HVAC_Inefficient'] = (df['HVAC_Energy'] > (20 + 10 * np.maximum(df['Temperature'] - 22, 0))).astype(int)

        df['Solar_Available'] = np.where((df['Hour'] >= 6) & (df['Hour'] <= 18), 100 * np.sin(np.pi * (df['Hour'] - 6) / 12), 0)
        df['Solar_Used'] = np.minimum(df['Solar_Available'], df['Total_Energy'])
        df['Grid_Energy'] = np.maximum(0, df['Total_Energy'] - df['Solar_Used'])
        df['CO2_Emissions'] = df['Grid_Energy'] * GRID_EMISSION_FACTOR + df['Solar_Used'] * SOLAR_EMISSION_FACTOR

        # Add lagged features
        df['Energy_Heavy_Lag1'] = df['Energy_Heavy'].shift(1).fillna(df['Energy_Heavy'].mean())
        df['Energy_Medium_Lag1'] = df['Energy_Medium'].shift(1).fillna(df['Energy_Medium'].mean())

        logger.info("Processed uploaded data with columns: %s", df.columns.tolist())
        return df
    except Exception as e:
        logger.error("Failed to process uploaded data: %s", e)
        st.error(f"Failed to process uploaded data: {e}")
        return None

def simulate_factory_data(num_heavy, energy_heavy, num_medium, energy_medium, p_ineff, q_ineff, r_ineff, hours=720):
    """Simulate factory energy consumption with inefficiencies and solar data."""
    try:
        timestamps = pd.date_range(start="2025-05-17 00:00", periods=hours, freq="H")
        df = pd.DataFrame({"Timestamp": timestamps})
        df["Day_of_Week"] = df["Timestamp"].dt.dayofweek
        df["Hour"] = df["Timestamp"].dt.hour
        df["Is_Weekday"] = df["Day_of_Week"] < 5

        conditions = [
            (~df["Is_Weekday"]),
            (df["Is_Weekday"] & (df["Hour"] >= 8) & (df["Hour"] < 16)),
            (df["Is_Weekday"] & (df["Hour"] >= 16) & (df["Hour"] < 24))
        ]
        choices = ["weekend", "day", "night"]
        df["Shift"] = np.select(conditions, choices, default="overnight")

        shift_map_heavy = {"day": num_heavy, "night": num_heavy // 2, "overnight": 0, "weekend": 0}
        shift_map_medium = {"day": num_medium, "night": num_medium, "overnight": num_medium // 5, "weekend": num_medium // 5}
        df["Intended_Heavy_On"] = df["Shift"].map(shift_map_heavy)
        df["Intended_Medium_On"] = df["Shift"].map(shift_map_medium)

        df["Heavy_On"] = df["Intended_Heavy_On"] + np.random.binomial(num_heavy - df["Intended_Heavy_On"], p_ineff)
        df["Medium_On"] = df["Intended_Medium_On"] + np.random.binomial(num_medium - df["Intended_Medium_On"], p_ineff)
        df["Energy_Heavy"] = df["Heavy_On"] * energy_heavy
        df["Energy_Medium"] = df["Medium_On"] * energy_medium

        df["Temperature"] = 30 + 5 * np.sin(2 * np.pi * (df["Hour"] - 14) / 24) + np.random.normal(0, 1, len(df))
        df["HVAC_Inefficient"] = np.random.choice([0, 1], size=len(df), p=[1 - q_ineff, q_ineff])
        df["HVAC_Energy"] = 20 + 10 * np.maximum(df["Temperature"] - (22 - 2 * df["HVAC_Inefficient"]), 0)

        df["Is_Working_Hours"] = df["Is_Weekday"] & (df["Hour"] >= 8) & (df["Hour"] < 18)
        df["Lighting_Energy"] = np.where(df["Is_Working_Hours"], 50, np.where(np.random.random(len(df)) < r_ineff, 50, 10))

        df["Solar_Available"] = np.where((df["Hour"] >= 6) & (df["Hour"] <= 18), 100 * np.sin(np.pi * (df["Hour"] - 6) / 12), 0)
        df["Solar_Used"] = np.minimum(df["Solar_Available"], df["Energy_Heavy"] + df["Energy_Medium"] + df["HVAC_Energy"] + df["Lighting_Energy"])
        df["Grid_Energy"] = np.maximum(0, df["Energy_Heavy"] + df["Energy_Medium"] + df["HVAC_Energy"] + df["Lighting_Energy"] - df["Solar_Used"])

        df["Total_Energy"] = df[["Energy_Heavy", "Energy_Medium", "HVAC_Energy", "Lighting_Energy"]].sum(axis=1)
        df["CO2_Emissions"] = df["Grid_Energy"] * GRID_EMISSION_FACTOR + df["Solar_Used"] * SOLAR_EMISSION_FACTOR

        df["Efficient_Heavy_On"] = df["Intended_Heavy_On"]
        df["Efficient_Medium_On"] = df["Intended_Medium_On"]
        df["Efficient_Energy_Heavy"] = df["Efficient_Heavy_On"] * energy_heavy
        df["Efficient_Energy_Medium"] = df["Efficient_Medium_On"] * energy_medium
        df["Efficient_HVAC_Energy"] = 20 + 10 * np.maximum(df["Temperature"] - 22, 0)
        df["Efficient_Lighting_Energy"] = np.where(df["Is_Working_Hours"], 50, 10)
        df["Efficient_Total_Energy"] = df[["Efficient_Energy_Heavy", "Efficient_Energy_Medium", "Efficient_HVAC_Energy", "Efficient_Lighting_Energy"]].sum(axis=1)
        df["Efficient_CO2_Emissions"] = df["Efficient_Total_Energy"] * GRID_EMISSION_FACTOR

        # Add lagged features
        df['Energy_Heavy_Lag1'] = df['Energy_Heavy'].shift(1).fillna(df['Energy_Heavy'].mean())
        df['Energy_Medium_Lag1'] = df['Energy_Medium'].shift(1).fillna(df['Energy_Medium'].mean())

        logger.info("Simulation completed with columns: %s", df.columns.tolist())
        return df
    except Exception as e:
        logger.error("Simulation failed: %s", e)
        st.error(f"Simulation failed: {e}")
        return None

class CarbonOptimizer:
    """Simple Q-learning for carbon-aware scheduling."""
    def __init__(self, actions=["shift_heavy", "shift_medium", "reduce_load", "no_action"], lr=0.1, gamma=0.9):
        self.q_table = {}
        self.actions = actions
        self.lr = lr
        self.gamma = gamma

    def get_state(self, shift, hour, solar_available):
        return (shift, hour, int(solar_available > 0))

    def get_action(self, state, epsilon=0.1):
        if np.random.random() < epsilon:
            return np.random.choice(self.actions)
        q_values = self.q_table.get(state, {a: 0 for a in self.actions})
        return max(q_values, key=q_values.get)

    def update(self, state, action, reward, next_state):
        if state not in self.q_table:
            self.q_table[state] = {a: 0 for a in self.actions}
        if next_state not in self.q_table:
            self.q_table[next_state] = {a: 0 for a in self.actions}
        current_q = self.q_table[state][action]
        next_max_q = max(self.q_table[next_state].values(), default=0)
        self.q_table[state][action] += self.lr * (reward + self.gamma * next_max_q - current_q)

@st.cache_data
def get_simulated_data(_num_heavy, _energy_heavy, _num_medium, _energy_medium, _p_ineff, _q_ineff, _r_ineff, _hours):
    return simulate_factory_data(_num_heavy, _energy_heavy, _num_medium, _energy_medium, _p_ineff, _q_ineff, _r_ineff, _hours)

@st.cache_resource
def train_model(df, target, model_type, cache_key):
    try:
        working_hours = df[df['Is_Working_Hours']]
        non_working_hours = df[~df['Is_Working_Hours']]

        features = ["Shift", "Hour", "Day_of_Week", "Temperature", "Is_Working_Hours", "Energy_Heavy_Lag1", "Energy_Medium_Lag1"]
        metrics = {}

        if not working_hours.empty:
            data = working_hours[features + [target]].dropna()
            X = data[features]
            y = data[target]

            X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.1, random_state=42)

            preprocessor = ColumnTransformer(
                transformers=[
                    ("shift", OneHotEncoder(categories=[['day', 'night', 'overnight', 'weekend']], drop="first", sparse_output=False), ["Shift"])
                ],
                remainder="passthrough"
            )
            if model_type == "Random Forest":
                model = RandomForestRegressor(n_estimators=100, max_depth=15, min_samples_split=5, random_state=42)
            else:
                model = LinearRegression()
            pipeline = Pipeline([("preprocessor", preprocessor), ("regressor", model)])
            pipeline.fit(X_train, y_train)
            y_pred = pipeline.predict(X_test)
            metrics["working"] = {
                "MAE": mean_absolute_error(y_test, y_pred),
                "MSE": mean_squared_error(y_test, y_pred),
                "R2": r2_score(y_test, y_pred)
            }
        else:
            metrics["working"] = {"MAE": 0, "MSE": 0, "R2": 0}

        if not non_working_hours.empty:
            y_pred_non_working = np.zeros(len(non_working_hours))
            y_true_non_working = non_working_hours[target]
            metrics["non_working"] = {
                "MAE": mean_absolute_error(y_true_non_working, y_pred_non_working),
                "MSE": mean_squared_error(y_true_non_working, y_pred_non_working),
                "R2": r2_score(y_true_non_working, y_pred_non_working) if y_true_non_working.var() > 0 else 0
            }
        else:
            metrics["non_working"] = {"MAE": 0, "MSE": 0, "R2": 0}

        full_data = df[features + [target]].dropna()
        X_full = full_data[features]
        y_full = full_data[target]
        predictions = np.zeros(len(X_full))
        working_indices = df[df['Is_Working_Hours']].index
        if not working_hours.empty:
            predictions[X_full.index.isin(working_indices)] = pipeline.predict(X_full.loc[working_indices])

        overall_metrics = {
            "MAE": mean_absolute_error(y_full, predictions),
            "MSE": mean_squared_error(y_full, predictions),
            "R2": r2_score(y_full, predictions)
        }
        overall_metrics["working"] = metrics["working"]
        overall_metrics["non_working"] = metrics["non_working"]

        joblib.dump(pipeline, f"model_{target}_{model_type.replace(' ', '_')}.joblib")
        return pipeline, overall_metrics
    except Exception as e:
        logger.error("Model training failed: %s", e)
        st.error(f"Model training failed: {e}")
        return None, None

def create_docx_from_markdown(markdown_text):
    """Convert markdown text to a .docx file."""
    try:
        doc = Document()
        lines = markdown_text.split('\n')
        for line in lines:
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('- '):
                doc.add_paragraph(line[2:], style='ListBullet')
            else:
                doc.add_paragraph(line)
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        return output
    except Exception as e:
        logger.error("Failed to create .docx: %s", e)
        st.error(f"Failed to create .docx: {e}")
        return None

def detect_inefficiencies(df, num_heavy, num_medium):
    """Detect inefficiencies in energy consumption data with dynamic thresholds."""
    inefficiencies = []
    
    # Machine inefficiencies during off-hours
    heavy_ineff = df[(df['Heavy_On'] > df['Intended_Heavy_On']) & (df['Shift'].isin(['overnight', 'weekend']))]
    medium_ineff = df[(df['Medium_On'] > df['Intended_Medium_On']) & (df['Shift'].isin(['overnight', 'weekend']))]
    if not heavy_ineff.empty:
        inefficiencies.append(f"Heavy machines running unnecessarily during {len(heavy_ineff)} off-hours (overnight/weekend), wasting {heavy_ineff['Energy_Heavy'].sum():.2f} kWh.")
    if not medium_ineff.empty:
        inefficiencies.append(f"Medium machines running unnecessarily during {len(medium_ineff)} off-hours (overnight/weekend), wasting {medium_ineff['Energy_Medium'].sum():.2f} kWh.")

    # HVAC inefficiencies
    hvac_ineff = df[df['HVAC_Inefficient'] == 1]
    if not hvac_ineff.empty:
        excess_hvac_energy = (hvac_ineff['HVAC_Energy'] - (20 + 10 * np.maximum(hvac_ineff['Temperature'] - 22, 0))).sum()
        inefficiencies.append(f"HVAC operating inefficiently for {len(hvac_ineff)} hours, wasting {excess_hvac_energy:.2f} kWh due to low temperature thresholds.")

    # Lighting inefficiencies
    lighting_ineff = df[(~df['Is_Working_Hours']) & (df['Lighting_Energy'] > 10.0)]
    if not lighting_ineff.empty:
        excess_lighting_energy = (lighting_ineff['Lighting_Energy'] - 10).sum()
        inefficiencies.append(f"Lighting left on unnecessarily during {len(lighting_ineff)} non-working hours, wasting {excess_lighting_energy:.2f} kWh.")

    # HVAC during low occupancy
    hvac_low_occupancy = df[(~df['Is_Working_Hours']) & (df['HVAC_Energy'] > 2.0)]
    if not hvac_low_occupancy.empty:
        excess_hvac_low = (hvac_low_occupancy['HVAC_Energy'] - 2).sum()
        inefficiencies.append(f"Excessive HVAC usage during {len(hvac_low_occupancy)} low-occupancy hours, wasting {excess_hvac_low:.2f} kWh.")

    # Dynamic inconsistency detection using statistical thresholds
    working_hours = df[df['Is_Working_Hours']]
    if not working_hours.empty:
        heavy_std = working_hours['Energy_Heavy'].std()
        heavy_avg = working_hours['Energy_Heavy'].mean()
        heavy_threshold = heavy_avg + 2 * heavy_std  # 2 standard deviations
        heavy_inconsistent = working_hours[working_hours['Energy_Heavy'] > heavy_threshold]
        if not heavy_inconsistent.empty:
            excess_heavy = (heavy_inconsistent['Energy_Heavy'] - heavy_avg).sum()
            inefficiencies.append(f"Inconsistent heavy machine usage during {len(heavy_inconsistent)} working hours (exceeding {heavy_threshold:.2f} kWh), wasting {excess_heavy:.2f} kWh.")

        medium_std = working_hours['Energy_Medium'].std()
        medium_avg = working_hours['Energy_Medium'].mean()
        medium_threshold = medium_avg + 2 * medium_std
        medium_inconsistent = working_hours[working_hours['Energy_Medium'] > medium_threshold]
        if not medium_inconsistent.empty:
            excess_medium = (medium_inconsistent['Energy_Medium'] - medium_avg).sum()
            inefficiencies.append(f"Inconsistent medium machine usage during {len(medium_inconsistent)} working hours (exceeding {medium_threshold:.2f} kWh), wasting {excess_medium:.2f} kWh.")

    # Solar underutilization with dynamic threshold
    solar_hours = df[(df['Hour'] >= 6) & (df['Hour'] <= 18) & (df['Solar_Available'] > 0)]
    if not solar_hours.empty:
        solar_avg_usage = solar_hours['Solar_Used'].mean()
        solar_underutilized = solar_hours[solar_hours['Solar_Used'] < 0.7 * solar_hours['Solar_Available']]  # 70% threshold
        if not solar_underutilized.empty:
            unused_solar = (solar_underutilized['Solar_Available'] - solar_underutilized['Solar_Used']).sum()
            inefficiencies.append(f"Solar energy underutilized during {len(solar_underutilized)} solar hours (usage <70% of available), missing {unused_solar:.2f} kWh.")

    return inefficiencies

def recommend_actions(inefficiencies, df):
    """Recommend actions based on detected inefficiencies and data patterns."""
    actions = []
    solar_hours = df[(df['Hour'] >= 6) & (df['Hour'] <= 18) & (df['Solar_Available'] > 0)]
    solar_usage_ratio = solar_hours['Solar_Used'].sum() / solar_hours['Solar_Available'].sum() if solar_hours['Solar_Available'].sum() > 0 else 0

    for ineff in inefficiencies:
        if "Heavy machines running unnecessarily" in ineff:
            actions.append("Implement automated shutdown protocols for heavy machines during overnight and weekend shifts to prevent unnecessary operation.")
        if "Medium machines running unnecessarily" in ineff:
            actions.append("Schedule medium machines to operate only during intended shifts (day/night) with IoT-based monitoring to enforce compliance.")
        if "HVAC operating inefficiently" in ineff:
            actions.append("Adjust HVAC systems to maintain a 22°C threshold during working hours and minimize overcooling, potentially saving significant energy.")
        if "Lighting left on unnecessarily" in ineff:
            actions.append("Deploy motion-sensor-based lighting controls to automatically turn off lights outside working hours (8 AM–6 PM weekdays).")
        if "Excessive HVAC usage" in ineff:
            actions.append("Configure HVAC to a minimal setting (e.g., 2 kWh) during non-working hours to reduce energy waste in low-occupancy periods.")
        if "Inconsistent heavy machine usage" in ineff:
            actions.append("Implement predictive maintenance and scheduling for heavy machines using real-time data to stabilize usage patterns.")
        if "Inconsistent medium machine usage" in ineff:
            actions.append("Optimize medium machine schedules with machine learning models to align with historical usage averages.")
        if "Solar energy underutilized" in ineff:
            actions.append(f"Shift {int((1 - solar_usage_ratio) * 100)}% of non-critical loads (e.g., medium machines) to solar hours (6 AM–6 PM) to maximize renewable energy usage.")

    # Additional recommendation based on solar usage
    if solar_usage_ratio < 0.7 and "Solar energy underutilized" not in [ineff.split(' (')[0] for ineff in inefficiencies]:
        actions.append(f"Increase solar energy utilization (current: {solar_usage_ratio*100:.1f}%) by scheduling high-energy tasks during peak solar hours (10 AM–2 PM).")

    return actions

def generate_dynamic_documentation(params, inefficiencies, actions, df):
    """Generate dynamic documentation based on simulation or uploaded data."""
    data_source = params.get('data_source', 'Simulated')
    num_heavy = params.get('num_heavy', 5)
    energy_heavy = params.get('energy_heavy', 20.0)
    num_medium = params.get('num_medium', 10)
    energy_medium = params.get('energy_medium', 10.0)
    cost_per_kwh = params.get('cost_per_kwh', 0.15)
    model_type = params.get('model_type', 'Random Forest')
    baseline_energy = df["Total_Energy"].sum()
    optimized_energy = df["Optimized_Total_Energy"].sum() if "Optimized_Total_Energy" in df.columns else baseline_energy
    energy_savings = baseline_energy - optimized_energy
    baseline_co2 = df["CO2_Emissions"].sum()
    optimized_co2 = df["Optimized_CO2_Emissions"].sum() if "Optimized_CO2_Emissions" in df.columns else baseline_co2
    co2_savings = baseline_co2 - optimized_co2
    trees_equivalent = (co2_savings / 1000) * TREES_PER_TON_CO2

    documentation = f"""
### Factory Energy & Carbon Optimizer Report

**Overview**  
This report details the energy and carbon optimization analysis for a factory, processed on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}. The analysis {'uses simulated data' if data_source == 'Simulate Data' else 'processes uploaded data'} to identify inefficiencies, recommend actions, and quantify savings. The tool supports UN Sustainable Development Goals (SDGs) 7 and 13 by reducing industrial energy waste and emissions.

**Configuration**  
- **Data Source**: {data_source}  
- **Heavy Machines**: {num_heavy} machines, each consuming {energy_heavy} kW  
- **Medium Machines**: {num_medium} machines, each consuming {energy_medium} kW  
- **Cost per kWh**: ${cost_per_kwh:.2f}  
- **Model Type**: {model_type}  
- **Data Period**: {df['Timestamp'].min().strftime('%Y-%m-%d')} to {df['Timestamp'].max().strftime('%Y-%m-%d')}  
- **Total Hours Analyzed**: {len(df)} hours  

**Detected Inefficiencies**  
{''.join([f'- {ineff}\n' for ineff in inefficiencies]) if inefficiencies else 'No significant inefficiencies detected.'}

**Recommended Actions**  
{''.join([f'- {action}\n' for action in actions]) if actions else 'No actions required.'}

**Results Summary**  
- **Baseline Energy**: {baseline_energy:.2f} kWh  
- **Optimized Energy**: {optimized_energy:.2f} kWh  
- **Energy Savings**: {energy_savings:.2f} kWh ({(energy_savings / baseline_energy * 100):.2f}% savings)  
- **Cost Savings**: ${energy_savings * cost_per_kwh:.2f}  
- **Baseline CO2 Emissions**: {baseline_co2:.2f} kg  
- **Optimized CO2 Emissions**: {optimized_co2:.2f} kg  
- **CO2 Savings**: {co2_savings:.2f} kg  
- **Trees Equivalent**: {trees_equivalent:.2f} trees  

**Technical Details**  
- **Simulation Assumptions**:  
  - Shifts: Day (8 AM–4 PM weekdays), Night (4 PM–12 AM weekdays), Overnight (12 AM–8 AM weekdays), Weekend (minimal operation).  
  - Emission Factors: Grid ({GRID_EMISSION_FACTOR} kg CO2/kWh), Solar ({SOLAR_EMISSION_FACTOR} kg CO2/kWh).  
  - Solar Availability: 100 kW peak from 6 AM–6 PM.  
  - Carbon Offset: {TREES_PER_TON_CO2} trees per ton of CO2 absorbed annually.  
- **Optimization**: Uses {model_type} for predicting machine schedules, combined with Q-learning for carbon-aware load shifting.  
- **Data Processing**: {'Synthetic data generated with configurable inefficiencies.' if data_source == 'Simulate Data' else 'Uploaded data validated and mapped to internal format.'}

**Conclusion**  
The analysis identified key inefficiencies and provided actionable recommendations to reduce energy consumption and emissions. By implementing the suggested actions, the factory can achieve significant cost and environmental benefits, aligning with sustainable industrial practices.
"""
    return documentation

def main():
    st.set_page_config(page_title="Factory Energy & Carbon Optimizer", layout="wide")
    st.title("Factory Energy & Carbon Optimizer")

    if st.sidebar.button("Clear Cache"):
        st.cache_data.clear()
        st.cache_resource.clear()
        st.rerun()

    tab1, tab2, tab3, tab4 = st.tabs(["Simulation", "Results", "Carbon Impact", "Documentation"])

    with tab1:
        st.header("Factory Setup")
        data_source = st.radio("Select Data Source", ["Simulate Data", "Upload Data"])

        if data_source == "Simulate Data":
            col1, col2 = st.columns(2)
            with col1:
                num_heavy = st.number_input("Number of Heavy Machines", min_value=1, value=5, step=1)
                energy_heavy = st.number_input("Energy per Heavy Machine (kW)", min_value=0.1, value=20.0, step=0.5)
                num_medium = st.number_input("Number of Medium Machines", min_value=1, value=10, step=1)
                energy_medium = st.number_input("Energy per Medium Machine (kW)", min_value=0.1, value=10.0, step=0.5)
            with col2:
                p_ineff = st.slider("Machine Inefficiency Probability", 0.0, 1.0, 0.1, 0.01)
                q_ineff = st.slider("HVAC Inefficiency Probability", 0.0, 1.0, 0.2, 0.01)
                r_ineff = st.slider("Lighting Inefficiency Probability", 0.0, 1.0, 0.1, 0.01)
                duration = st.selectbox("Simulation Duration", ["1 Week (168h)", "1 Month (720h)", "3 Months (2160h)"], index=1)
                hours = {"1 Week (168h)": 168, "1 Month (720h)": 720, "3 Months (2160h)": 2160}[duration]
        else:
            uploaded_file = st.file_uploader("Upload Energy Data (Excel/CSV)", type=['xlsx', 'csv'])
            num_heavy = 5
            energy_heavy = 20.0
            num_medium = 10
            energy_medium = 10.0
            p_ineff = q_ineff = r_ineff = 0.1
            hours = None

        cost_per_kwh = st.number_input("Cost per kWh ($)", min_value=0.0, value=0.15, step=0.01)
        model_type = st.selectbox("Prediction Model", ["Random Forest", "Linear Regression"])

        if st.button("Run"):
            with st.spinner("Processing..."):
                if data_source == "Simulate Data":
                    df = get_simulated_data(num_heavy, energy_heavy, num_medium, energy_medium, p_ineff, q_ineff, r_ineff, hours)
                else:
                    if uploaded_file is None:
                        st.error("Please upload a file.")
                        return
                    df = process_uploaded_data(uploaded_file)

                if df is not None:
                    st.session_state["df"] = df
                    st.session_state["params"] = {
                        "num_heavy": num_heavy,
                        "energy_heavy": energy_heavy,
                        "num_medium": num_medium,
                        "energy_medium": energy_medium,
                        "cost_per_kwh": cost_per_kwh,
                        "model_type": model_type,
                        "data_source": data_source
                    }
                    st.success("Data processed successfully!")
                    st.rerun()

    with tab2:
        if "df" in st.session_state and "params" in st.session_state:
            df = st.session_state["df"]
            params = st.session_state["params"]
            st.header("Optimization Results")
            required_cols = ["Shift", "Hour", "Solar_Available", "CO2_Emissions", "Intended_Heavy_On", "Intended_Medium_On",
                           "Day_of_Week", "Heavy_On", "Medium_On", "Energy_Heavy", "Energy_Medium", "HVAC_Energy",
                           "Lighting_Energy", "Temperature", "HVAC_Inefficient", "Is_Working_Hours", "Solar_Used", "Grid_Energy"]
            if not all(col in df.columns for col in required_cols):
                st.error("Missing required columns. Please re-run the simulation or upload valid data.")
            else:
                try:
                    st.subheader("Detected Inefficiencies")
                    inefficiencies = detect_inefficiencies(df, params["num_heavy"], params["num_medium"])
                    if inefficiencies:
                        for ineff in inefficiencies:
                            st.write(f"- {ineff}")
                    else:
                        st.write("No significant inefficiencies detected.")

                    st.subheader("Recommended Actions")
                    actions = recommend_actions(inefficiencies, df)
                    if actions:
                        for action in actions:
                            st.write(f"- {action}")
                    else:
                        st.write("No actions required.")

                    with st.spinner("Optimizing..."):
                        heavy_model, heavy_metrics = train_model(df, "Intended_Heavy_On", params["model_type"], "heavy")
                        medium_model, medium_metrics = train_model(df, "Intended_Medium_On", params["model_type"], "medium")

                    if heavy_model and medium_model:
                        st.subheader("Model Accuracy")
                        st.markdown("**Heavy Machines Model**")
                        st.table({
                            "Metric": ["Mean Absolute Error (MAE)", "Mean Squared Error (MSE)", "R-squared (R²)",
                                       "MAE (Working Hours)", "R² (Working Hours)", "MAE (Non-Working Hours)", "R² (Non-Working Hours)"],
                            "Value": [f"{heavy_metrics['MAE']:.4f}", f"{heavy_metrics['MSE']:.4f}", f"{heavy_metrics['R2']:.4f}",
                                      f"{heavy_metrics['working']['MAE']:.4f}", f"{heavy_metrics['working']['R2']:.4f}",
                                      f"{heavy_metrics['non_working']['MAE']:.4f}", f"{heavy_metrics['non_working']['R2']:.4f}"]
                        })
                        st.markdown("**Medium Machines Model**")
                        st.table({
                            "Metric": ["Mean Absolute Error (MAE)", "Mean Squared Error (MSE)", "R-squared (R²)",
                                       "MAE (Working Hours)", "R² (Working Hours)", "MAE (Non-Working Hours)", "R² (Non-Working Hours)"],
                            "Value": [f"{medium_metrics['MAE']:.4f}", f"{medium_metrics['MSE']:.4f}", f"{medium_metrics['R2']:.4f}",
                                      f"{medium_metrics['working']['MAE']:.4f}", f"{medium_metrics['working']['R2']:.4f}",
                                      f"{medium_metrics['non_working']['MAE']:.4f}", f"{medium_metrics['non_working']['R2']:.4f}"]
                        })

                        optimizer = CarbonOptimizer()
                        for i in range(min(100, len(df) - 1)):
                            state = optimizer.get_state(df["Shift"].iloc[i], df["Hour"].iloc[i], df["Solar_Available"].iloc[i])
                            action = optimizer.get_action(state)
                            reward = -df["CO2_Emissions"].iloc[i] / (df["Total_Energy"].iloc[i] + 1)
                            next_state = optimizer.get_state(df["Shift"].iloc[i + 1], df["Hour"].iloc[i + 1], df["Solar_Available"].iloc[i + 1])
                            optimizer.update(state, action, reward, next_state)

                        df["Predicted_Heavy_On"] = heavy_model.predict(df[["Shift", "Hour", "Day_of_Week", "Temperature", "Is_Working_Hours", "Energy_Heavy_Lag1", "Energy_Medium_Lag1"]]).round().clip(0, params["num_heavy"])
                        df["Predicted_Medium_On"] = medium_model.predict(df[["Shift", "Hour", "Day_of_Week", "Temperature", "Is_Working_Hours", "Energy_Heavy_Lag1", "Energy_Medium_Lag1"]]).round().clip(0, params["num_medium"])
                        df["Optimized_Heavy_On"] = df["Intended_Heavy_On"].copy()
                        df["Optimized_Medium_On"] = df["Intended_Medium_On"].copy()

                        for i in range(len(df)):
                            state = optimizer.get_state(df["Shift"].iloc[i], df["Hour"].iloc[i], df["Solar_Available"].iloc[i])
                            action = optimizer.get_action(state, epsilon=0.1)
                            if action == "shift_heavy" and df["Solar_Available"].iloc[i] > 0 and df["Hour"].iloc[i] < 18:
                                df.loc[i, "Optimized_Heavy_On"] = min(df["Predicted_Heavy_On"].iloc[i], params["num_heavy"])
                            elif action == "shift_medium" and df["Solar_Available"].iloc[i] > 0 and df["Hour"].iloc[i] < 18:
                                df.loc[i, "Optimized_Medium_On"] = min(df["Predicted_Medium_On"].iloc[i], params["num_medium"])
                            elif action == "reduce_load" and df["Hour"].iloc[i] >= 18:
                                df.loc[i, "Optimized_Heavy_On"] = max(0, df["Optimized_Heavy_On"].iloc[i] - 1)
                                df.loc[i, "Optimized_Medium_On"] = max(0, df["Optimized_Medium_On"].iloc[i] - 1)

                        df["Optimized_Energy_Heavy"] = df["Optimized_Heavy_On"] * params["energy_heavy"]
                        df["Optimized_Energy_Medium"] = df["Optimized_Medium_On"] * params["energy_medium"]
                        df["Optimized_HVAC_Energy"] = np.where(
                            df["HVAC_Inefficient"] == 0, df["HVAC_Energy"], 20 + 10 * np.maximum(df["Temperature"] - 22, 0)
                        )
                        df["Optimized_Lighting_Energy"] = np.where(df["Is_Working_Hours"] | (df["Lighting_Energy"] <= 10.0), df["Lighting_Energy"], 10.0)
                        df["Optimized_Total_Energy"] = df[["Optimized_Energy_Heavy", "Optimized_Energy_Medium", "Optimized_HVAC_Energy", "Optimized_Lighting_Energy"]].sum(axis=1)

                        df["Optimized_Solar_Used"] = np.minimum(df["Solar_Available"], df["Optimized_Total_Energy"])
                        df["Optimized_Grid_Energy"] = np.maximum(0, df["Optimized_Total_Energy"] - df["Optimized_Solar_Used"])
                        df["Optimized_CO2_Emissions"] = df["Optimized_Grid_Energy"] * GRID_EMISSION_FACTOR + df["Optimized_Solar_Used"] * SOLAR_EMISSION_FACTOR

                        df["Optimized_Total_Energy"] = np.where(df["Optimized_Total_Energy"] > df["Total_Energy"] * MAX_ENERGY_INCREASE,
                                                              df["Total_Energy"] * MAX_ENERGY_INCREASE, df["Optimized_Total_Energy"])
                        df["Optimized_Grid_Energy"] = np.maximum(0, df["Optimized_Total_Energy"] - df["Optimized_Solar_Used"])
                        df["Optimized_CO2_Emissions"] = df["Optimized_Grid_Energy"] * GRID_EMISSION_FACTOR + df["Optimized_Solar_Used"] * SOLAR_EMISSION_FACTOR

                        baseline_energy = df["Total_Energy"].sum()
                        optimized_energy = df["Optimized_Total_Energy"].sum()
                        energy_savings = baseline_energy - optimized_energy
                        savings_percent = (energy_savings / baseline_energy * 100) if baseline_energy > 0 else 0
                        cost_savings = energy_savings * params["cost_per_kwh"]
                        baseline_co2 = df["CO2_Emissions"].sum()
                        optimized_co2 = df["Optimized_CO2_Emissions"].sum()
                        co2_savings = baseline_co2 - optimized_co2
                        trees_equivalent = (co2_savings / 1000) * TREES_PER_TON_CO2

                        st.subheader("Summary")
                        st.table({
                            "Metric": ["Baseline Energy (kWh)", "Optimized Energy (kWh)", "Energy Savings (kWh)", "Savings (%)",
                                       "Cost Savings ($)", "Baseline CO2 (kg)", "Optimized CO2 (kg)", "CO2 Savings (kg)",
                                       "Trees Equivalent"],
                            "Value": [f"{baseline_energy:.2f}", f"{optimized_energy:.2f}", f"{energy_savings:.2f}",
                                      f"{savings_percent:.2f}", f"{cost_savings:.2f}", f"{baseline_co2:.2f}",
                                      f"{optimized_co2:.2f}", f"{co2_savings:.2f}", f"{trees_equivalent:.2f}"]
                        })

                        st.subheader("Energy Usage")
                        fig_energy = go.Figure()
                        fig_energy.add_trace(go.Scatter(x=df["Timestamp"], y=df["Total_Energy"], name="Baseline", line=dict(color="red")))
                        fig_energy.add_trace(go.Scatter(x=df["Timestamp"], y=df["Optimized_Total_Energy"], name="Optimized", line=dict(color="green")))
                        fig_energy.update_layout(
                            title="Energy Usage (kW)",
                            yaxis_title="Energy (kW)",
                            xaxis_title="Timestamp",
                            hovermode="x unified",
                            legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1)
                        )
                        st.plotly_chart(fig_energy, use_container_width=True)

                        st.subheader("Energy Breakdown")
                        fig_breakdown = go.Figure()
                        for col, name, color in [("Energy_Heavy", "Heavy Machines", "blue"), ("Energy_Medium", "Medium Machines", "orange"),
                                               ("HVAC_Energy", "HVAC", "green"), ("Lighting_Energy", "Lighting", "purple")]:
                            fig_breakdown.add_trace(go.Scatter(x=df["Timestamp"], y=df[col], name=name, stackgroup="one", line=dict(color=color)))
                        fig_breakdown.update_layout(
                            title="Baseline Energy Breakdown (kW)",
                            yaxis_title="Energy (kW)",
                            xaxis_title="Timestamp",
                            hovermode="x unified",
                            legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1),
                            showlegend=True
                        )
                        st.plotly_chart(fig_breakdown, use_container_width=True)

                        st.subheader("Daily Savings")
                        df["Date"] = df["Timestamp"].dt.date
                        daily_savings = df.groupby("Date").apply(lambda x: x["Total_Energy"].sum() - x["Optimized_Total_Energy"].sum()).reset_index(name="Savings")
                        daily_savings = daily_savings[daily_savings["Savings"] != 0]
                        fig_savings = px.bar(daily_savings, x="Date", y="Savings", title="Daily Energy Savings (kWh)", color_discrete_sequence=["teal"])
                        fig_savings.update_layout(
                            yaxis_title="Savings (kWh)",
                            xaxis_title="Date",
                            hovermode="x unified",
                            showlegend=False
                        )
                        st.plotly_chart(fig_savings, use_container_width=True)

                        st.subheader("Export")
                        csv = df.to_csv(index=False)
                        st.download_button("Download Data as CSV", csv, f"factory_energy_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv", "text/csv")

                except Exception as e:
                    st.error(f"Optimization error: {e}")
                    logger.error(f"Optimization error: {e}")

    with tab3:
        if "df" in st.session_state:
            df = st.session_state["df"]
            st.header("Carbon Impact")
            if "CO2_Emissions" not in df.columns or "Optimized_CO2_Emissions" not in df.columns:
                st.error("Carbon data missing. Please re-run the simulation or upload valid data.")
            else:
                try:
                    st.subheader("Emissions Over Time")
                    fig_co2_time = go.Figure()
                    fig_co2_time.add_trace(go.Scatter(x=df["Timestamp"], y=df["CO2_Emissions"], name="Baseline", line=dict(color="red")))
                    fig_co2_time.add_trace(go.Scatter(x=df["Timestamp"], y=df["Optimized_CO2_Emissions"], name="Optimized", line=dict(color="green")))
                    fig_co2_time.update_layout(
                        title="CO2 Emissions Over Time (kg)",
                        yaxis_title="CO2 Emissions (kg)",
                        xaxis_title="Timestamp",
                        hovermode="x unified",
                        legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1)
                    )
                    st.plotly_chart(fig_co2_time, use_container_width=True)

                    st.subheader("Emissions Breakdown")
                    co2_breakdown = {
                        "Heavy Machines": (df["Energy_Heavy"] * GRID_EMISSION_FACTOR).sum(),
                        "Medium Machines": (df["Energy_Medium"] * GRID_EMISSION_FACTOR).sum(),
                        "HVAC": (df["HVAC_Energy"] * GRID_EMISSION_FACTOR).sum(),
                        "Lighting": (df["Lighting_Energy"] * GRID_EMISSION_FACTOR).sum()
                    }
                    fig_pie = px.pie(names=list(co2_breakdown.keys()), values=list(co2_breakdown.values()), title="Baseline CO2 Breakdown")
                    fig_pie.update_layout(
                        hovermode="x unified",
                        showlegend=True
                    )
                    st.plotly_chart(fig_pie, use_container_width=True)

                    st.subheader("Daily CO2 Emissions Comparison")
                    df["Date"] = df["Timestamp"].dt.date
                    daily_co2 = df.groupby("Date").agg({
                        "CO2_Emissions": "sum",
                        "Optimized_CO2_Emissions": "sum"
                    }).reset_index()
                    daily_co2 = daily_co2[(daily_co2["CO2_Emissions"] > 0) | (daily_co2["Optimized_CO2_Emissions"] > 0)]
                    fig_daily_co2 = go.Figure()
                    fig_daily_co2.add_trace(go.Bar(
                        x=daily_co2["Date"],
                        y=daily_co2["CO2_Emissions"],
                        name="Baseline",
                        marker_color="red"
                    ))
                    fig_daily_co2.add_trace(go.Bar(
                        x=daily_co2["Date"],
                        y=daily_co2["Optimized_CO2_Emissions"],
                        name="Optimized",
                        marker_color="green"
                    ))
                    fig_daily_co2.update_layout(
                        title="Daily CO2 Emissions Comparison (kg)",
                        yaxis_title="CO2 Emissions (kg)",
                        xaxis_title="Date",
                        barmode="group",
                        hovermode="x unified",
                        legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1)
                    )
                    st.plotly_chart(fig_daily_co2, use_container_width=True)

                    baseline_co2 = df["CO2_Emissions"].sum()
                    optimized_co2 = df["Optimized_CO2_Emissions"].sum()
                    co2_savings = baseline_co2 - optimized_co2
                    trees_equivalent = (co2_savings / 1000) * TREES_PER_TON_CO2

                    st.subheader("Sustainability Impact")
                    st.markdown(f"""
                    - **CO2 Savings**: {co2_savings:.2f} kg
                    - **Trees Equivalent**: {trees_equivalent:.2f} trees
                    - **Goal**: Supports UN SDG 7 & 13 by reducing industrial emissions.
                    """)

                except Exception as e:
                    st.error(f"Carbon impact error: {e}")
                    logger.error(f"Carbon impact error: {e}")

    with tab4:
        st.header("Documentation")
        if "df" in st.session_state and "params" in st.session_state:
            df = st.session_state["df"]
            params = st.session_state["params"]
            inefficiencies = detect_inefficiencies(df, params["num_heavy"], params["num_medium"])
            actions = recommend_actions(inefficiencies, df)
            documentation_text = generate_dynamic_documentation(params, inefficiencies, actions, df)
        else:
            documentation_text = """
            ### Factory Energy & Carbon Optimizer
            **Overview**  
            Please run a simulation or upload data to generate a detailed report.
            """
        st.markdown(documentation_text)

        st.subheader("Export Documentation")
        docx_file = create_docx_from_markdown(documentation_text)
        if docx_file:
            st.download_button(
                label="Download Documentation as DOCX",
                data=docx_file,
                file_name=f"factory_energy_documentation_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

if __name__ == "__main__":
    main()
